"""Unit Tests for Document Processing, Parsers, Cleaners, and Metadata Extraction.

Covers:
- TextCleaner Unicode normalization, control character removal, whitespace collapsing, token stats
- MetadataParser PDF date extraction, ISO dates, YAML frontmatter parsing
- PdfParser PyMuPDF extraction, multi-page streaming, metadata, and corrupt PDF handling
- DocxParser python-docx paragraph/table extraction, core properties, and corrupt DOCX handling
- MarkdownParser YAML frontmatter extraction, heading inference, and text body parsing
- TextParser multi-encoding support (UTF-8, UTF-8-BOM, Latin-1, CSV, JSON, Logs)
- DocumentParserFactory format resolution, registry management, and unsupported type rejection
- DocumentProcessor asynchronous and synchronous end-to-end processing pipelines
"""

import io
import uuid
from datetime import UTC, datetime

import docx
import fitz  # PyMuPDF
import pytest

from app.document_processing import (
    CorruptedDocumentException,
    DocumentParserFactory,
    DocumentProcessor,
    DocxParser,
    MarkdownParser,
    MetadataParser,
    PdfParser,
    ProcessingResult,
    TextCleaner,
    TextParser,
    UnsupportedDocumentException,
)

# ------------------------------------------------------------------------------
# 1. Text Cleaner Tests
# ------------------------------------------------------------------------------


def test_text_cleaner_unicode_normalization() -> None:
    """Test NFKC Unicode normalization for ligatures and compatibility characters."""
    # Ligatures: ﬁ -> fi, ﬂ -> fl
    ligature_text = "The \ufb01le is on the \ufb02oor."
    normalized = TextCleaner.normalize_unicode(ligature_text)
    assert normalized == "The file is on the floor."

    # Fullwidth characters: Investiga
    fullwidth_text = "\uff29\uff4e\uff56\uff45\uff53\uff54\uff49\uff47\uff41"
    assert TextCleaner.normalize_unicode(fullwidth_text) == "Investiga"


def test_text_cleaner_control_characters_and_null_bytes() -> None:
    """Test stripping of null bytes, invisible zero-width chars, and control characters."""
    dirty_text = "Clean\x00Text\u200bWith\ufeffZero\u00adWidth\x07\x0bChars"
    cleaned = TextCleaner.remove_control_characters(dirty_text)
    assert cleaned == "CleanTextWithZeroWidthChars"


def test_text_cleaner_line_endings_and_whitespace_collapsing() -> None:
    """Test deterministic line endings and structural paragraph preservation."""
    crlf_text = "Paragraph One with   excessive    spaces.\r\n\r\n\r\n\r\nParagraph Two.\r\nLine 2.1."
    cleaned = TextCleaner.clean(crlf_text)

    expected = "Paragraph One with excessive spaces.\n\nParagraph Two.\nLine 2.1."
    assert cleaned == expected


def test_text_cleaner_word_and_char_counts() -> None:
    """Test accurate statistical counting for words and characters."""
    sample = "Investiga Incident Response Platform for Security Operations"
    assert TextCleaner.count_words(sample) == 7
    assert TextCleaner.count_characters(sample) == len(sample)
    assert TextCleaner.count_words("") == 0
    assert TextCleaner.count_characters("") == 0


# ------------------------------------------------------------------------------
# 2. Metadata Parser Tests
# ------------------------------------------------------------------------------


def test_metadata_parser_pdf_dates() -> None:
    """Test parsing of standard PDF date strings."""
    # UTC format
    dt1 = MetadataParser.parse_pdf_date("D:20260805123045Z")
    assert dt1 == datetime(2026, 8, 5, 12, 30, 45, tzinfo=UTC)

    # Offset format (+02'00')
    dt2 = MetadataParser.parse_pdf_date("D:20260805140000+02'00'")
    assert dt2 is not None
    assert dt2.year == 2026 and dt2.month == 8 and dt2.day == 5

    # Invalid string returns None
    assert MetadataParser.parse_pdf_date("NotADate") is None
    assert MetadataParser.parse_pdf_date(None) is None


def test_metadata_parser_iso_dates() -> None:
    """Test parsing of standard ISO-8601 timestamps."""
    dt = MetadataParser.parse_iso_date("2026-08-05T18:00:00Z")
    assert dt == datetime(2026, 8, 5, 18, 0, 0, tzinfo=UTC)

    assert MetadataParser.parse_iso_date("invalid-date") is None


def test_metadata_parser_yaml_frontmatter() -> None:
    """Test extracting YAML frontmatter from document header."""
    markdown_sample = (
        "---\n"
        "title: Kubernetes Incident Runbook\n"
        "author: SRE Lead\n"
        "date: 2026-08-05\n"
        "category: Runbook\n"
        "---\n"
        "# Incident Triage Procedures\n"
        "Follow these steps..."
    )
    meta, body = MetadataParser.extract_frontmatter(markdown_sample)
    assert meta["title"] == "Kubernetes Incident Runbook"
    assert meta["author"] == "SRE Lead"
    assert meta["category"] == "Runbook"
    assert body.startswith("# Incident Triage Procedures")

    # Document without frontmatter
    no_fm, no_fm_body = MetadataParser.extract_frontmatter("Just plain text.")
    assert no_fm == {}
    assert no_fm_body == "Just plain text."


# ------------------------------------------------------------------------------
# 3. PDF Parser Tests
# ------------------------------------------------------------------------------


def test_pdf_parser_text_and_metadata_extraction() -> None:
    """Test PDF text extraction and metadata retrieval from multi-page document."""
    # Build in-memory multi-page PDF using PyMuPDF
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((50, 50), "Page 1: System Architecture Overview")

    page2 = doc.new_page()
    page2.insert_text((50, 50), "Page 2: Database Migration Strategy")

    doc.set_metadata(
        {
            "title": "Investiga Technical Spec",
            "author": "Chief Architect",
            "creationDate": "D:20260805120000Z",
        }
    )
    pdf_bytes = doc.tobytes()
    doc.close()

    parser = PdfParser()
    assert parser.supports(".pdf") is True
    assert parser.supports(".PDF") is True
    assert parser.supports(".txt") is False

    extracted = parser.parse(pdf_bytes)
    assert "Page 1: System Architecture Overview" in extracted.raw_text
    assert "Page 2: Database Migration Strategy" in extracted.raw_text
    assert extracted.metadata.page_count == 2
    assert extracted.metadata.title == "Investiga Technical Spec"
    assert extracted.metadata.author == "Chief Architect"
    assert extracted.metadata.creation_date is not None


def test_pdf_parser_corrupted_payload() -> None:
    """Test corrupt PDF byte stream handling."""
    parser = PdfParser()
    with pytest.raises(CorruptedDocumentException):
        parser.parse(b"This is not a valid PDF file stream %PDF-invalid")


# ------------------------------------------------------------------------------
# 4. DOCX Parser Tests
# ------------------------------------------------------------------------------


def test_docx_parser_paragraphs_tables_and_metadata() -> None:
    """Test DOCX paragraph, table, and Core Properties extraction."""
    # Build in-memory DOCX
    doc = docx.Document()
    doc.core_properties.title = "Incident Report 2026"
    doc.core_properties.author = "Security Analyst"

    doc.add_paragraph("Summary: Denial of Service attack detected on edge gateway.")
    doc.add_paragraph("Remediation: Traffic rate limiting activated.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Dropped Packets"
    table.cell(1, 1).text = "1,500,000"

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    parser = DocxParser()
    assert parser.supports(".docx") is True
    assert parser.supports(".dotx") is True
    assert parser.supports(".pdf") is False

    extracted = parser.parse(docx_bytes)
    assert "Summary: Denial of Service attack detected" in extracted.raw_text
    assert "Remediation: Traffic rate limiting activated." in extracted.raw_text
    assert "Metric | Value" in extracted.raw_text
    assert "Dropped Packets | 1,500,000" in extracted.raw_text
    assert extracted.metadata.title == "Incident Report 2026"
    assert extracted.metadata.author == "Security Analyst"


def test_docx_parser_corrupted_payload() -> None:
    """Test corrupted DOCX payload handling."""
    parser = DocxParser()
    with pytest.raises(CorruptedDocumentException):
        parser.parse(b"Not a zip file or valid docx package")


# ------------------------------------------------------------------------------
# 5. Markdown Parser Tests
# ------------------------------------------------------------------------------


def test_markdown_parser_with_frontmatter_and_headings() -> None:
    """Test Markdown parsing with frontmatter and H1 title resolution."""
    sample_md = (
        "---\n"
        "title: Root Cause Analysis\n"
        "author: SRE Team\n"
        "lang: en\n"
        "---\n"
        "# Incident #4092 Postmortem\n\n"
        "Root cause was identified as memory leak in worker pool."
    )

    parser = MarkdownParser()
    assert parser.supports(".md") is True
    assert parser.supports(".markdown") is True

    extracted = parser.parse(sample_md.encode("utf-8"))
    assert extracted.metadata.title == "Root Cause Analysis"
    assert extracted.metadata.author == "SRE Team"
    assert extracted.metadata.language == "en"
    assert "Root cause was identified as memory leak" in extracted.raw_text


def test_markdown_parser_without_frontmatter_heading_fallback() -> None:
    """Test Markdown title extraction from level 1 heading when frontmatter is absent."""
    sample_md = "# Disaster Recovery Drill 2026\n\nDrill conducted successfully."
    parser = MarkdownParser()
    extracted = parser.parse(sample_md.encode("utf-8"))
    assert extracted.metadata.title == "Disaster Recovery Drill 2026"
    assert "Drill conducted successfully." in extracted.raw_text


# ------------------------------------------------------------------------------
# 6. Text Parser Tests
# ------------------------------------------------------------------------------


def test_text_parser_encodings_and_formats() -> None:
    """Test Plaintext parser with various text formats and encodings."""
    parser = TextParser()
    assert parser.supports(".txt") is True
    assert parser.supports(".log") is True
    assert parser.supports(".csv") is True
    assert parser.supports(".json") is True
    assert parser.supports(".yaml") is True

    # UTF-8 with BOM
    bom_content = "\ufeff2026-08-05 ERROR [auth] Failed login attempt".encode(
        "utf-8-sig"
    )
    extracted = parser.parse(bom_content)
    assert "Failed login attempt" in extracted.raw_text

    # Latin-1 encoded characters
    latin1_content = "Café résumé".encode("latin-1")
    extracted_latin = parser.parse(latin1_content)
    assert "Café" in extracted_latin.raw_text or "Caf" in extracted_latin.raw_text


# ------------------------------------------------------------------------------
# 7. Document Parser Factory Tests
# ------------------------------------------------------------------------------


def test_parser_factory_resolution_and_unsupported_formats() -> None:
    """Test factory resolution of parsers and rejection of unsupported formats."""
    factory = DocumentParserFactory()

    assert isinstance(factory.get_parser("report.pdf"), PdfParser)
    assert isinstance(factory.get_parser(".pdf"), PdfParser)
    assert isinstance(factory.get_parser("guide.docx"), DocxParser)
    assert isinstance(factory.get_parser("README.md"), MarkdownParser)
    assert isinstance(factory.get_parser("server.log"), TextParser)
    assert isinstance(factory.get_parser("data.csv"), TextParser)
    assert isinstance(factory.get_parser("config.yaml"), TextParser)

    # Rejection of unsupported extension
    with pytest.raises(UnsupportedDocumentException):
        factory.get_parser("payload.exe")

    with pytest.raises(UnsupportedDocumentException):
        factory.get_parser("archive.tar.gz")


# ------------------------------------------------------------------------------
# 8. Document Processor End-to-End Tests
# ------------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_document_processor_async_end_to_end() -> None:
    """Test DocumentProcessor async execution pipeline with cleaning and statistics."""
    processor = DocumentProcessor()
    doc_id = uuid.uuid4()

    dirty_md = (
        "---\n"
        "title: Security Audit\n"
        "author: SecOps\n"
        "---\n"
        "# Security Audit Findings\r\n\r\n"
        "Finding 1:   Unauthenticated   endpoint   detected.\x00\u200b\r\n\r\n\r\n\r\n"
        "Action: Fixed in release 1.2."
    )

    result: ProcessingResult = await processor.process(
        content=dirty_md.encode("utf-8"),
        filename="audit_report.md",
        document_id=doc_id,
        language="en",
    )

    assert result.document_id == doc_id
    assert result.title == "Security Audit"
    assert result.author == "SecOps"
    assert result.language == "en"
    assert result.page_count == 1
    assert result.word_count > 0
    assert result.character_count > 0
    assert result.processing_time_ms >= 0.0

    # Ensure clean text has no null bytes, zero-width chars, or CRLF
    assert "\x00" not in result.clean_text
    assert "\u200b" not in result.clean_text
    assert "\r" not in result.clean_text
    assert "Finding 1: Unauthenticated endpoint detected." in result.clean_text


def test_document_processor_sync_execution() -> None:
    """Test DocumentProcessor synchronous execution convenience method."""
    processor = DocumentProcessor()
    raw_content = b"Simple sync text processing test line."

    result = processor.process_sync(
        content=raw_content,
        filename="test.txt",
    )

    assert result.clean_text == "Simple sync text processing test line."
    assert result.word_count == 6
    assert result.title == "test"
    assert result.processing_time_ms >= 0.0
