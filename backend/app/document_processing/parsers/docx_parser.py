"""DOCX Document Parser using python-docx.

Implements text and structured table content extraction, along with Core Properties
metadata retrieval for Microsoft Word (.docx) files.
"""

import io
from datetime import UTC
from pathlib import Path
from typing import ClassVar

import docx
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.document_processing.exceptions import CorruptedDocumentException
from app.document_processing.models import ExtractedDocument, ExtractedMetadata
from app.document_processing.parsers.base_parser import BaseDocumentParser


class DocxParser(BaseDocumentParser):
    """Parser for Microsoft Word OpenXML (.docx) documents."""

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {".docx", ".dotx"}
    SUPPORTED_MIME_TYPES: ClassVar[set[str]] = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.template",
        "application/msword",
    }

    def supports(self, extension: str, mime_type: str | None = None) -> bool:
        """Evaluate if parser supports DOCX format."""
        norm_ext = (
            extension.lower() if extension.startswith(".") else f".{extension.lower()}"
        )
        if norm_ext in self.SUPPORTED_EXTENSIONS:
            return True
        if mime_type and mime_type.lower() in self.SUPPORTED_MIME_TYPES:
            return True
        return False

    def _open_doc(self, content: bytes | Path) -> DocxDocument:
        """Safely open python-docx Document object."""
        try:
            if isinstance(content, (str, Path)):
                return docx.Document(str(content))
            return docx.Document(io.BytesIO(content))
        except (PackageNotFoundError, Exception) as exc:
            filename = (
                str(content) if isinstance(content, (str, Path)) else "in-memory.docx"
            )
            raise CorruptedDocumentException(
                filename=filename,
                reason=f"Failed to open DOCX package: {exc}",
            ) from exc

    def extract_text(self, content: bytes | Path) -> str:
        """Extract text from paragraphs and tables."""
        doc = self._open_doc(content)
        text_elements: list[str] = []

        # 1. Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text and text.strip():
                text_elements.append(text)

        # 2. Extract structured table cells
        for table in doc.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_cells:
                    text_elements.append(" | ".join(row_cells))

        return "\n\n".join(text_elements)

    def extract_metadata(self, content: bytes | Path) -> ExtractedMetadata:
        """Extract Core Properties from DOCX package."""
        doc = self._open_doc(content)
        core_props = doc.core_properties

        title = (
            core_props.title if core_props.title and core_props.title.strip() else None
        )
        author = (
            core_props.author
            if core_props.author and core_props.author.strip()
            else None
        )

        creation_date = core_props.created
        if creation_date and creation_date.tzinfo is None:
            creation_date = creation_date.replace(tzinfo=UTC)

        modification_date = core_props.modified
        if modification_date and modification_date.tzinfo is None:
            modification_date = modification_date.replace(tzinfo=UTC)

        return ExtractedMetadata(
            title=title,
            author=author,
            creation_date=creation_date,
            modification_date=modification_date,
            page_count=1,
            extra_metadata={
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
                "last_modified_by": core_props.last_modified_by,
            },
        )

    def parse(self, content: bytes | Path) -> ExtractedDocument:
        """Execute single-pass text and metadata extraction."""
        doc = self._open_doc(content)

        text_elements: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text and text.strip():
                text_elements.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_cells:
                    text_elements.append(" | ".join(row_cells))

        raw_text = "\n\n".join(text_elements)

        core_props = doc.core_properties
        title = (
            core_props.title if core_props.title and core_props.title.strip() else None
        )
        author = (
            core_props.author
            if core_props.author and core_props.author.strip()
            else None
        )

        creation_date = core_props.created
        if creation_date and creation_date.tzinfo is None:
            creation_date = creation_date.replace(tzinfo=UTC)

        modification_date = core_props.modified
        if modification_date and modification_date.tzinfo is None:
            modification_date = modification_date.replace(tzinfo=UTC)

        metadata = ExtractedMetadata(
            title=title,
            author=author,
            creation_date=creation_date,
            modification_date=modification_date,
            page_count=1,
            extra_metadata={
                "subject": core_props.subject,
                "category": core_props.category,
                "last_modified_by": core_props.last_modified_by,
            },
        )

        return ExtractedDocument(raw_text=raw_text, metadata=metadata)
